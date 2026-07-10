import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF
from datetime import datetime

class PDFReport(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, 'MediCast - Autonomous Surgical Intelligence Report', 0, 1, 'C')
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def generate_docx_report(context: list[dict], session_id: str = "Unknown") -> io.BytesIO:
    doc = Document()
    
    # Title
    title = doc.add_heading('MediCast Autonomous Surgical Intelligence Report', 0)
    title.alignment = 1 # center
    
    # Metadata
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Session ID: {session_id}")
    doc.add_paragraph(f"Total Frames Analyzed: {len(context)}")
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        "This report contains an AI-driven analysis of a surgical procedure. "
        "MediCast autonomously analyzes surgical video streams to detect anatomical landmarks, "
        "evaluate safety protocols, track procedural phases, and provide actionable educational feedback."
    )
    
    critical_events = [f for f in context if f.get('arbiter', {}).get('escalationLevel') == 'CRITICAL']
    warnings = [f for f in context if f.get('arbiter', {}).get('escalationLevel') == 'WARNING']
    
    doc.add_paragraph(f"Critical Events Detected: {len(critical_events)}", style='List Bullet')
    doc.add_paragraph(f"Warnings Detected: {len(warnings)}", style='List Bullet')
    
    doc.add_heading('Detailed Event Log', level=1)
    
    for frame in context:
        frame_index = frame.get('frameIndex', 0)
        arbiter = frame.get('arbiter', {})
        escalation = arbiter.get('escalationLevel', 'SAFE')
        summary = arbiter.get('summary', 'No summary available.')
        action = arbiter.get('recommendedAction', 'None')
        
        # Add heading for frame
        p = doc.add_paragraph()
        run = p.add_run(f"Frame {frame_index} - [{escalation}]")
        run.bold = True
        if escalation == 'CRITICAL':
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif escalation == 'WARNING':
            run.font.color.rgb = RGBColor(255, 165, 0)
        else:
            run.font.color.rgb = RGBColor(0, 128, 0)
            
        doc.add_paragraph(f"Observation: {summary}")
        doc.add_paragraph(f"Recommended Action: {action}")
        
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_pdf_report(context: list[dict], session_id: str = "Unknown") -> io.BytesIO:
    pdf = PDFReport()
    pdf.set_margins(10, 10, 10)
    pdf.add_page()
    
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 1)
    pdf.cell(0, 10, f"Session ID: {session_id}", 0, 1)
    pdf.cell(0, 10, f"Total Frames Analyzed: {len(context)}", 0, 1)
    
    critical_events = [f for f in context if f.get('arbiter', {}).get('escalationLevel') == 'CRITICAL']
    warnings = [f for f in context if f.get('arbiter', {}).get('escalationLevel') == 'WARNING']
    
    pdf.ln(5)
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(190, 10, 'Executive Summary', 0, 1)
    pdf.set_font('Arial', '', 11)
    pdf.multi_cell(190, 8, (
        "This report contains an AI-driven analysis of a surgical procedure. "
        "MediCast autonomously analyzes surgical video streams to detect anatomical landmarks, "
        "evaluate safety protocols, track procedural phases, and provide actionable educational feedback."
    ))
    
    pdf.ln(5)
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 8, f"Critical Events Detected: {len(critical_events)}", 0, 1)
    pdf.cell(0, 8, f"Warnings Detected: {len(warnings)}", 0, 1)
    
    pdf.ln(10)
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, 'Detailed Event Log', 0, 1)
    
    for frame in context:
        frame_index = frame.get('frameIndex', 0)
        arbiter = frame.get('arbiter', {})
        escalation = arbiter.get('escalationLevel', 'SAFE')
        summary = arbiter.get('summary', 'No summary available.')
        action = arbiter.get('recommendedAction', 'None')
        
        pdf.set_font('Arial', 'B', 11)
        if escalation == 'CRITICAL':
            pdf.set_text_color(200, 0, 0)
        elif escalation == 'WARNING':
            pdf.set_text_color(200, 100, 0)
        else:
            pdf.set_text_color(0, 150, 0)
            
        # Clean special chars that fpdf might trip on
        summary = summary.encode('latin-1', 'replace').decode('latin-1')
        action = action.encode('latin-1', 'replace').decode('latin-1')
        
        pdf.cell(190, 8, f"Frame {frame_index} - [{escalation}]", 0, 1)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(190, 6, f"Observation: {summary}")
        pdf.multi_cell(190, 6, f"Recommended Action: {action}")
        pdf.ln(4)
        
    file_stream = io.BytesIO()
    pdf_bytes = pdf.output()
    file_stream.write(bytes(pdf_bytes))
    file_stream.seek(0)
    return file_stream
